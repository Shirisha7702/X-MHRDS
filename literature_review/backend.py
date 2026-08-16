import os
import urllib.request
import urllib.parse
import xml.etree.ElementTree as ET
import ssl
import time
import zipfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from literature_review.metadata import PAPERS_METADATA

def search_arxiv_by_title(title):
    """
    Search arXiv API by title to find a matching preprint and retrieve its PDF URL.
    """
    print(f"  Attempting to search arXiv for preprint of: '{title}'...")
    query = f'ti:"{title}"'
    params = {
        "search_query": query,
        "max_results": 1
    }
    query_string = urllib.parse.urlencode(params, quote_via=urllib.parse.quote_plus)
    url = f"http://export.arxiv.org/api/query?{query_string}"
    
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    
    try:
        # Be polite to arXiv QPS limits
        time.sleep(3.0)
        with urllib.request.urlopen(req, context=ctx) as response:
            xml_data = response.read()
            root = ET.fromstring(xml_data)
            ns = {"atom": "http://www.w3.org/2005/Atom"}
            entry = root.find("atom:entry", ns)
            if entry is not None:
                pdf_url = ""
                for link_node in entry.findall("atom:link", ns):
                    if link_node.get("title") == "pdf" or link_node.get("type") == "application/pdf":
                        pdf_url = link_node.get("href")
                if pdf_url:
                    print(f"  Found preprint PDF on arXiv: {pdf_url}")
                    return pdf_url
    except Exception as e:
        print(f"  Error searching arXiv: {e}")
    return None

def download_pdf(url, output_path):
    """
    Download a PDF file from a URL.
    """
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    try:
        with urllib.request.urlopen(req, context=ctx) as response:
            with open(output_path, "wb") as f:
                f.write(response.read())
        print(f"  Successfully downloaded PDF to: {output_path}")
        return True
    except Exception as e:
        print(f"  Failed to download PDF from {url}: {e}")
        return False

def write_placeholder(paper, output_dir):
    """
    Write a text metadata placeholder for paywalled or unavailable papers.
    """
    filename = f"{paper['id']}_placeholder.txt"
    filepath = os.path.join(output_dir, filename)
    content = f"""METADATA PLACEHOLDER FOR PAYWALLED PAPER

Title: {paper.get('title')}
Authors: {paper.get('authors')}
Year: {paper.get('year')}
URL / DOI: {paper.get('url')}

---
MANUAL RETRIEVAL INSTRUCTIONS:
This paper is behind a publisher paywall and was not found as a free preprint on arXiv.
To manually retrieve the full text:
1. Navigate to the URL: {paper.get('url')}
2. Access the PDF via institutional login or subscribing credentials.
3. Place the downloaded PDF file into this folder and rename it to '{paper['id']}.pdf' to update the archive.
"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"  Created metadata placeholder: {filepath}")

def create_docx(output_path):
    """
    Generate a styled Microsoft Word document using python-docx.
    """
    doc = Document()
    
    # Styles Setup
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Title Cover Page styling
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("\n\n\n\n\n\n\nEXPLAINABLE MENTAL HEALTH RISK DETECTION SYSTEM\n")
    run_title.bold = True
    run_title.font.size = Pt(24)
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_subtitle = p_subtitle.add_run("Automated Academic Literature Review and Synthesis (2024 - 2026)\n")
    run_subtitle.italic = True
    run_subtitle.font.size = Pt(14)
    run_subtitle.font.color.rgb = RGBColor(0x5C, 0x76, 0x8D) # Muted blue
    
    p_details = doc.add_paragraph()
    p_details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_details = p_details.add_run(
        "\n\n\n\n\n\nPrepared for: Advanced Artificial Intelligence Projects in Data Science (55-710603)\n"
        "Created by: Shirisha Srirangam, Vara Prasad Kurella, Sai Krishna Samudrapu, John Babu Thammisetti, Raviteja Vibhuthi\n"
        "Date: July 2026\n"
    )
    run_details.font.size = Pt(11)
    run_details.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_page_break()
    
    # Executive Summary Table
    h1_summary = doc.add_paragraph()
    run_h1_sum = h1_summary.add_run("1. Executive Summary Table")
    run_h1_sum.bold = True
    run_h1_sum.font.size = Pt(18)
    run_h1_sum.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    doc.add_paragraph(
        "This table aggregates the 16 primary peer-reviewed papers and preprints identified in this literature review, "
        "providing a high-level reference of their key focus areas and metadata."
    )
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Paper ID'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Authors'
    hdr_cells[3].text = 'Year'
    
    for paper in PAPERS_METADATA:
        row_cells = table.add_row().cells
        row_cells[0].text = paper['id']
        row_cells[1].text = paper['title']
        row_cells[2].text = paper['authors'].split(',')[0] + " et al." if ',' in paper['authors'] else paper['authors']
        row_cells[3].text = paper['year']
        
    doc.add_page_break()
    
    # Detailed Literature Reviews
    h1_reviews = doc.add_paragraph()
    run_h1_rev = h1_reviews.add_run("2. Detailed Literature Reviews")
    run_h1_rev.bold = True
    run_h1_rev.font.size = Pt(18)
    run_h1_rev.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    for idx, paper in enumerate(PAPERS_METADATA, 1):
        h2 = doc.add_paragraph()
        run_h2 = h2.add_run(f"2.{idx} {paper['title']}")
        run_h2.bold = True
        run_h2.font.size = Pt(14)
        run_h2.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88)
        
        p_bib = doc.add_paragraph()
        run_bib_lbl = p_bib.add_run("Citation: ")
        run_bib_lbl.bold = True
        p_bib.add_run(paper['citation'])
        
        p_meth = doc.add_paragraph()
        run_meth_lbl = p_meth.add_run("Core Methodology: ")
        run_meth_lbl.bold = True
        p_meth.add_run(paper['methodology'])
        
        p_grid_lbl = doc.add_paragraph()
        run_grid_lbl = p_grid_lbl.add_run("Critique Summary:")
        run_grid_lbl.bold = True
        
        grid_table = doc.add_table(rows=1, cols=2)
        grid_table.style = 'Light Shading Accent 1'
        grid_hdr = grid_table.rows[0].cells
        grid_hdr[0].text = "Pros / Strengths"
        grid_hdr[1].text = "Cons / Limitations"
        
        grid_cells = grid_table.add_row().cells
        
        pros_text = ""
        for p in paper['pros']:
            pros_text += f"\u2022 {p}\n"
        grid_cells[0].text = pros_text.strip()
        
        cons_text = ""
        for c in paper['cons']:
            cons_text += f"\u2022 {c}\n"
        grid_cells[1].text = cons_text.strip()
        
        doc.add_paragraph("")
        
    doc.save(output_path)
    print(f"Generated Word document at: {output_path}")

def make_zip(workspace_root, zip_path):
    """
    Compress the literature_review.docx, literature_review.md,
    and all contents of papers/ folder into a zip file at workspace root.
    """
    review_dir = os.path.join(workspace_root, "literature_review")
    docx_path = os.path.join(review_dir, "literature_review.docx")
    md_path = os.path.join(review_dir, "literature_review.md")
    papers_dir = os.path.join(review_dir, "papers")
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        if os.path.exists(docx_path):
            zipf.write(docx_path, arcname="literature_review.docx")
        if os.path.exists(md_path):
            zipf.write(md_path, arcname="literature_review.md")
            
        if os.path.exists(papers_dir):
            for foldername, subfolders, filenames in os.walk(papers_dir):
                for filename in filenames:
                    filepath = os.path.join(foldername, filename)
                    rel_path = os.path.relpath(filepath, papers_dir)
                    zipf.write(filepath, arcname=os.path.join("papers", rel_path))
                    
    print(f"Successfully created zip archive at: {zip_path}")
