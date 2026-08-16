"""
Converts this project's Markdown documentation (PROJECT_PLAYBOOK.md, DEMO_GUIDE.md) into
properly formatted Word .docx files: real Heading styles (so Word's Navigation Pane and an
auto-generated Table of Contents both work), real tables, real bullet/numbered lists, code
spans/blocks in a monospace font, and clickable hyperlinks -- external URLs, and internal
"[text](#anchor)" cross-references resolved to real Word bookmarks so they jump to the right
heading inside the document rather than rendering as dead text.

Usage: python scripts/md_to_docx.py <input.md> <output.docx> "<Document Title>" "<Subtitle>"
"""
import re
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Calibri"
CODE_FONT = "Consolas"
BRAND_COLOR = RGBColor(0x0B, 0x4F, 0x8A)
CODE_BG = "F0F0F0"
HEADER_ROW_BG = "1F4E79"
LINK_COLOR = RGBColor(0x0B, 0x4F, 0x8A)


def slugify(text):
    """Mirrors GitHub's heading-anchor algorithm so '#some-heading' links in the source
    Markdown resolve to the same heading a browser/VS Code preview would jump to."""
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]*)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]*)\*", r"\1", text)
    text = text.lower()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"\s", "-", text)
    return text


class DocxBuilder:
    def __init__(self, title, subtitle):
        self.doc = Document()
        self._setup_styles()
        self._bookmark_id = 0
        self._slug_to_bookmark = {}
        self._pending_bookmarks = []  # slugs waiting to be attached to the next heading
        self._build_title_page(title, subtitle)
        self._build_toc_page()

    # ------------------------------------------------------------------ styles

    def _setup_styles(self):
        doc = self.doc
        normal = doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

        for lvl, size, color in [(1, 20, BRAND_COLOR), (2, 15, BRAND_COLOR), (3, 12.5, RGBColor(0x22, 0x22, 0x22))]:
            h = doc.styles[f"Heading {lvl}"]
            h.font.name = BODY_FONT
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = color
            h.paragraph_format.space_before = Pt(18 if lvl == 1 else 12)
            h.paragraph_format.space_after = Pt(8 if lvl == 1 else 6)
            h.paragraph_format.page_break_before = (lvl == 1)

        if "Code" not in [s.name for s in doc.styles]:
            code_style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
            code_style.base_style = doc.styles["Normal"]
            code_style.font.name = CODE_FONT
            code_style.font.size = Pt(9.5)
            code_style.paragraph_format.space_before = Pt(2)
            code_style.paragraph_format.space_after = Pt(2)
            code_style.paragraph_format.left_indent = Inches(0.15)

        if "Quote" not in [s.name for s in doc.styles]:
            quote_style = doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
            quote_style.base_style = doc.styles["Normal"]
            quote_style.font.italic = True
            quote_style.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            quote_style.paragraph_format.left_indent = Inches(0.3)
            quote_style.paragraph_format.space_before = Pt(6)
            quote_style.paragraph_format.space_after = Pt(6)

    # ------------------------------------------------------------------ title / TOC

    def _build_title_page(self, title, subtitle):
        doc = self.doc
        for _ in range(4):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = BRAND_COLOR

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run("X-MHRDS — Explainable Mental Health Risk Detection System")
        run3.font.size = Pt(11)
        run3.font.italic = True
        run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        self._add_page_break()

    def _build_toc_page(self):
        doc = self.doc
        h = doc.add_paragraph("Table of Contents")
        h.style = doc.styles["Heading 1"]
        h.paragraph_format.page_break_before = False

        note = doc.add_paragraph()
        note.add_run(
            "(Auto-generated field — in Word, right-click it and choose \"Update Field\" "
            "→ \"Update entire table\" if it shows as empty when you first open this file.)"
        ).italic = True

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_text = OxmlElement("w:t")
        fld_text.text = "Right-click and select Update Field to generate the table of contents."
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_text)
        run._r.append(fld_end)

        self._add_page_break()

    def _add_page_break(self):
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)

    # ------------------------------------------------------------------ bookmarks / links

    def _next_bookmark_name(self):
        self._bookmark_id += 1
        return f"bm{self._bookmark_id}"

    def register_heading_slug(self, slug):
        """Called while pre-scanning the document so every internal link target is known
        before we render (a link to a heading further down still needs to resolve)."""
        if slug not in self._slug_to_bookmark:
            self._slug_to_bookmark[slug] = self._next_bookmark_name()
        return self._slug_to_bookmark[slug]

    def _insert_bookmark(self, paragraph, name):
        bm_id = str(abs(hash(name)) % 1000000)
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bm_id)
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bm_id)
        paragraph._p.insert(0, start)
        paragraph._p.append(end)

    def _add_hyperlink(self, paragraph, text, url_or_anchor, is_internal, code=False):
        hyperlink = OxmlElement("w:hyperlink")
        if is_internal:
            bookmark_name = self._slug_to_bookmark.get(url_or_anchor)
            if bookmark_name is None:
                # Unresolvable internal link (rare) -- fall back to plain styled text.
                self._add_run(paragraph, text, bold=False, italic=False, code=code)
                return
            hyperlink.set(qn("w:anchor"), bookmark_name)
        else:
            r_id = paragraph.part.relate_to(
                url_or_anchor,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0B4F8A")
        rpr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rpr.append(underline)
        if code:
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), CODE_FONT)
            rfonts.set(qn("w:hAnsi"), CODE_FONT)
            rpr.append(rfonts)
        new_run.append(rpr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    # ------------------------------------------------------------------ inline formatting

    INLINE_TOKEN_RE = re.compile(
        r"(?P<code>`[^`]+`)"
        r"|(?P<link>\[[^\]]+\]\([^)]+\))"
        r"|(?P<bold>\*\*[^*]+\*\*)"
        r"|(?P<italic>\*[^*]+\*)"
    )

    def _add_run(self, paragraph, text, bold=False, italic=False, code=False):
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
        return run

    def add_inline(self, paragraph, text):
        """Parses **bold**, *italic*, `code`, and [label](url) within a line and appends
        correctly-formatted runs (including real hyperlinks) to the given paragraph."""
        pos = 0
        for m in self.INLINE_TOKEN_RE.finditer(text):
            if m.start() > pos:
                self._add_run(paragraph, text[pos:m.start()])
            if m.group("code"):
                self._add_run(paragraph, m.group("code")[1:-1], code=True)
            elif m.group("link"):
                label, target = re.match(r"\[([^\]]+)\]\(([^)]+)\)", m.group("link")).groups()
                is_code_label = label.startswith("`") and label.endswith("`")
                clean_label = label[1:-1] if is_code_label else label
                if target.startswith("#"):
                    self._add_hyperlink(paragraph, clean_label, target[1:], True, code=is_code_label)
                else:
                    self._add_hyperlink(paragraph, clean_label, target, False, code=is_code_label)
            elif m.group("bold"):
                self._add_run(paragraph, m.group("bold")[2:-2], bold=True)
            elif m.group("italic"):
                self._add_run(paragraph, m.group("italic")[1:-1], italic=True)
            pos = m.end()
        if pos < len(text):
            self._add_run(paragraph, text[pos:])

    # ------------------------------------------------------------------ block builders

    def add_heading(self, level, text, slug):
        level = min(level, 3)
        p = self.doc.add_paragraph(style=f"Heading {level}")
        self.add_inline(p, text)
        bookmark_name = self._slug_to_bookmark.get(slug)
        if bookmark_name:
            self._insert_bookmark(p, bookmark_name)
        return p

    def add_paragraph(self, text):
        p = self.doc.add_paragraph()
        self.add_inline(p, text)
        return p

    def add_quote(self, text):
        p = self.doc.add_paragraph(style="Quote")
        self.add_inline(p, text)
        return p

    def add_bullet(self, text, ordered=False, level=0):
        style = "List Number" if ordered else "List Bullet"
        p = self.doc.add_paragraph(style=style)
        if level:
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        self.add_inline(p, text)
        return p

    def add_code_block(self, lines):
        for i, line in enumerate(lines):
            p = self.doc.add_paragraph(style="Code")
            if i == 0:
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), CODE_BG)
                pPr.append(shd)
            self._shade_paragraph(p)
            run = p.add_run(line if line.strip() else " ")
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)

    def _shade_paragraph(self, p):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), CODE_BG)
        pPr.append(shd)

    def add_hr(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_table(self, header, rows):
        table = self.doc.add_table(rows=1, cols=len(header))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.autofit = True

        hdr_cells = table.rows[0].cells
        for i, col_text in enumerate(header):
            cell = hdr_cells[i]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            self.add_inline(p, col_text)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), HEADER_ROW_BG)
            tcPr.append(shd)

        for r_idx, row in enumerate(rows):
            cells = table.add_row().cells
            for i, cell_text in enumerate(row):
                if i >= len(cells):
                    continue
                cell = cells[i]
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                self.add_inline(p, cell_text)
                if r_idx % 2 == 1:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), "F4F7FB")
                    tcPr.append(shd)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_footer_page_numbers(self):
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    def save(self, path):
        self.add_footer_page_numbers()
        self.doc.save(path)


# ------------------------------------------------------------------ markdown parsing

def parse_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line):
    return bool(re.match(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?\s*$", line))


def parse_markdown(lines):
    """Yields (block_type, payload) tuples for the subset of Markdown actually used in
    this project's docs: headings, hr, blockquote, fenced code, tables, bullet/numbered
    lists, and plain paragraphs (including the multi-line Table-of-Contents link list)."""
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            yield ("code", code_lines)
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            yield ("heading", (len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        if re.match(r"^-{3,}$", stripped):
            yield ("hr", None)
            i += 1
            continue

        if stripped.startswith(">"):
            yield ("quote", stripped.lstrip(">").strip())
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < n and is_table_separator(lines[i + 1]):
            header = parse_table_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i].strip()))
                i += 1
            yield ("table", (header, rows))
            continue

        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            level = len(m.group(1)) // 2
            yield ("bullet", (m.group(2), level))
            i += 1
            continue

        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            level = len(m.group(1)) // 2
            yield ("number", (m.group(2), level))
            i += 1
            continue

        # Plain paragraph -- accumulate contiguous non-blank, non-special lines.
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#{1,6}\s|-{3,}$|>|\||```|\s*[-*]\s|\s*\d+\.\s)", lines[i]
        ):
            para_lines.append(lines[i].strip())
            i += 1
        yield ("para", " ".join(para_lines))


def convert(md_path, docx_path, title, subtitle):
    with open(md_path, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()

    blocks = list(parse_markdown(raw_lines))

    builder = DocxBuilder(title, subtitle)

    # Pass 1: pre-register every heading's slug -> bookmark so forward-references resolve.
    for block_type, payload in blocks:
        if block_type == "heading":
            _, text = payload
            builder.register_heading_slug(slugify(text))

    # Pass 2: render.
    for block_type, payload in blocks:
        if block_type == "heading":
            level, text = payload
            builder.add_heading(level, text, slugify(text))
        elif block_type == "para":
            builder.add_paragraph(payload)
        elif block_type == "quote":
            builder.add_quote(payload)
        elif block_type == "hr":
            builder.add_hr()
        elif block_type == "code":
            builder.add_code_block(payload)
        elif block_type == "table":
            header, rows = payload
            builder.add_table(header, rows)
        elif block_type == "bullet":
            text, level = payload
            builder.add_bullet(text, ordered=False, level=level)
        elif block_type == "number":
            text, level = payload
            builder.add_bullet(text, ordered=True, level=level)

    builder.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 5:
        print("Usage: python md_to_docx.py <input.md> <output.docx> <title> <subtitle>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4])
